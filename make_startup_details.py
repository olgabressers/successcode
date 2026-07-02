from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

GOLD = RGBColor(0xB8, 0x88, 0x2A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Georgia'; r.font.size = Pt(20)
    r.font.color.rgb = GOLD; r.bold = True

def subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(text)
    r.font.size = Pt(10); r.font.color.rgb = GREY; r.italic = True

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B8882A')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.color.rgb = GOLD; r.font.size = Pt(13)

def label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = GREY

def body(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(10); r.bold = bold

# ── TITLE ─────────────────────────────────────────────────────────
title('SuccessCode — Startup Details')
subtitle('a16z Speedrun Application · 2026')

# ── 1. RELEVANT EXPERIENCE ────────────────────────────────────────
section_heading('1.  Relevant Experience')
label('Form question:')
body('In a few sentences, highlight your most relevant professional, startup, or industry experience. Focus on track record, domain expertise, and past wins.')
label('Answer:')
body(
    '20 years as a Director at Richemont, Adidas, Nike, IBM — running €150M digital P&L, leading '
    'e-commerce & digital transformation across Europe. Simultaneously, 10 years of deep study in '
    'Zi Wei Dou Shu Imperial Science under leading Asian masters, applied personally to every '
    'high-stakes career decision, negotiation, and transition since 2016. Tested methodology with '
    '600+ clients. The product is built on a decade of live validation before a single line of code '
    'was written. No other founder in this space has operated at both the execution scale required '
    'to build a platform and the methodological depth required to encode the system correctly.'
)

rule()

# ── 2. ONE-SENTENCE PITCH ─────────────────────────────────────────
section_heading('2.  One-Sentence Pitch')
label('Form question:')
body('Pitch your startup in one sentence. What do you do and for whom?')
label('Answer (9 words):')
body('AI decision and timing intelligence for everyone, science based.', bold=True)

rule()

# ── 3. STARTUP DESCRIPTION ────────────────────────────────────────
section_heading('3.  Startup Description')
label('Form question:')
body('What problem are you solving? What are you building?')
label('Answer (100 words):')
body(
    'Every person faces high-stakes decisions — career, partnerships, money, health, investments — '
    'with no personal decision or timing intelligence layer. SuccessCode solves this.'
)
body(
    'Zi Wei Dou Shu is a 3,000-year-old mathematically derived system — used by emperors, built '
    'from precise birth data — that maps individual timing windows, risk periods, and winning angles. '
    'Not generic horoscopes. Structural personal intelligence, used by Eastern businesses for centuries.'
)
body(
    'The calculator is live at successcode.net. Phase 2 builds the AI model answering user questions '
    'and providing annual, decade, lifetime trends. Phase 3 targets 500+ subscribers by Q4 2026.'
)
body('No product yet offers this precision. SuccessCode is that product.')

rule()

# ── 4. TEAM ───────────────────────────────────────────────────────
section_heading('4.  Team')
label('Form question:')
body('Tell us more about the team. How does the team know each other? Is there anyone else on the team? Why is this the best team to win? Do you have any key advisors?')
label('Answer (82 words):')
body(
    'Currently solo. The product is built on 10 years of specialized Zi Wei Dou Shu study — '
    'the methodology and domain expertise are my core competitive advantage. Phase 1 was built '
    'fully solo using AI development tools, and I am capable of building the complete product '
    'independently. I will bring on a technical co-founder or senior engineer as the product '
    'scales into Phase 2. Startup advisors will be engaged at that stage. This is a deliberate '
    'sequencing choice: domain knowledge had to come first. The execution infrastructure comes next.'
)

rule()

# ── 5. LINKS ──────────────────────────────────────────────────────
section_heading('5.  Links')
label('LinkedIn URL:')
body('https://www.linkedin.com/in/olgabressers/')
label('GitHub URL:')
body('https://github.com/olgabressers  (repo is private — code protection)')
label('X URL:')
body('N/A')
label('Portfolio URL:')
body('https://successcode.net/founderportfolio/')
label('Personal website / project:')
body('https://successcode.net/speedrun/')

# ── SAVE ──────────────────────────────────────────────────────────
out = "/Users/olgaigoshina/Desktop/Documents/Private/CLAUDE AI/Zi Wei Dou Shu/SuccessCode/Startup Details.docx"
doc.save(out)
print(f"Saved: {out}")
import os
print(f"Size: {os.path.getsize(out):,} bytes")
